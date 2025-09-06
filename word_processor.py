import docx 
import os
from dataclass.word_document_content import WordDocumentContent 

class WordProcessor: 
    def process_docx(self, file_path):
        """
        Process a .docx file and extract text content
        Args: file_path: Path to the .docx file

        Returns:
        WordDocumentContent: Extracted text content
        dict: Metadata about the document"""
    



        try:  
            doc = docx. Document (file_path) 
            # Extract text from paragraphs BANNED)
            paragraphs = [] 
            for para in doc. paragraphs: 
                if para.text.strip():
                    paragraphs .append(para. text) 

            #extract text from tables 
            table_texts = []
            for table in doc.tables:    
                table_text = [] 
                for row in table.rows:
                    row_text = [] 
                    for cell in row.cells:     
                        if cell.text.strip(): 
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text)) 
                if table_text:
                    table_texts.append("\n". join(table_text))

            # Combine paragraph and table text  
            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n" + "\n\n".join(table_texts)

            # Create document content
            # content = WordDocumentContent(full text)

            # Create metadata
            metadata = {

                "page count": len(doc.sections),
                "has_tables": len(doc.tables) > 0,
                "paragraph_count": len(doc.paragraphs),
                "file_size" : os.path.getsize(file_path) 
            }

            return full_text, metadata
        except Exception as e :
            raise Exception(f"Error processing Word document: {e}") 

    def process_doc(self, file_path):
        """
        Process a doc file and extract text content
        Args:
            file_path: Path to the .doc file
        Returns:
            wordDocumentContent: Extracted text content
            dict: Metadata about the document
        """ 
        # Create content object
        content = WordDocumentContent()
        metadata = {
            "page_count": 1, # Default
            "file_size": os.path.getsize(file_path) 
         }
         # Try multiple methods to extract text from .doc files 
        extraction_methods = [
            self._extract_with_pywin32,
            self._extract_with_textract, 
            self._extract_with_win32com, 
            self._extract_with_docx2txt, 
            self._extract_with_olefile
        ]

        last_error = None 
        for method in extraction_methods:
            try:
                text = method(file_path)
                if text and len(text) > 10: # Ensure we got some meaningful content ig
                    content = WordDocumentContent(text) 
                    return content, metadata 
            except Exception as e:
                last_error = e
                continue # Try next method a

        # Tf we got here, all methods Failed
        if last_error:
            raise Exception(f"Failed to extract text from doc file after trying all methods: (last error)")
        else:
            #Last resort return empty content with a note about failure
            content = WordDocumentContent("[Could not extract text from this doc file. The document may be password-protected, corrupt, or using a format that is not supported")
            return content, metadata


    def _extract_with_textract(self, file_path):
        """Try to extract text using textract"""
        import textract
        try:
            text = textract.process(file_path, method='antiword').decode("utf-8")
        except:
            #If antiword fails, try other textract methods
            text = textract.process(file_path).decode("utf-8")
        return text

    def _extract_with_win32com(self, file_path):
        """Try to extract text using win32com with Word application"""
        try:
            import win32com.client
            import os

            word = win32com.client.Dispatch("Word.Application")
            word.visible = False

            #Convert to absolute path
            file_path=os.path.abspath(file_path)

            try:
                doc = word. Documents.Open(file_path, False, True, False)
                text = doc.Content.Text
                doc.Close(False)
                word.Quit()
                return text
            except Exception as e:
                if word:
                    word.Quit()   
                raise e
        except ImportError:
            raise Exception("win32com not available")

    def _extract_with_docx2txt(self, file_path):
        """Try to extract text using docx2txt"""
        import docx2txt
        return docx2txt.process(file_path)
    
    def _extract_with_pywin32(self, file_path):
        """Another method using pywin32"""
        try:    
            from win32com import client
            import pythoncom

            pythoncom.Colnitialize()
            word_app = client.Dispatch("Word.Application")
            word_app.Visible = False    

            doc = word_app.Documents.Open(file_path)
            text = ""
            for para in doc. Paragraphs:
                text += para.Range.Text +"\in"
            doc.Close()
            word_app.Quit()
            pythoncom.Couninitialize()
            return text
        except ImportError:
            raise Exception("pywin32 not available")
    def _extract_with_olefile(self, file_path):
        """Extract text using olefile for very old .doc files"""
        try:
            import olefile
            if olefile.isoleFile(file_path):
                with olefile. OleFile(file_path) as ole:
                    if ole.exists("Wordflucument"):
                        #Just extract text from WordDocument stream
                        #This is a simplified approach and won't get all text
                        word_stream = ole.openstream("WordDocument")
                        content = word_stream.read().decode('utf-8', errors = 'ignore')

                        #Basic cleanup of binary content
                        import re
                        #Keep only printable ASCII characters
                        content = re.sub(r'[^\x20\x7E\r\n]','',content)
                        #Remove sequences of whitespace
                        content = re.sub(r'\s+','', content)

                        return content
            raise Exception("Not a valid OLE file")
        except ImportError:

            raise Exception("olefile not available")

if __name__ == "main":
    word_processor = WordProcessor()
    word_processor.process_docx(r'file_path')
